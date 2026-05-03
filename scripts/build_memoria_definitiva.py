from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt


BASE = Path(r"c:\Users\Sebas\Desktop\py3")
SOURCE_DOCX = BASE / "Memoria_TFG_IES_El_Grao_v2_con_wireframes_mockups.docx"
SOURCE_PDF = BASE / "Memoria_Juan_Sebastián_Agudelo_DAW.pdf"
OUTPUT_DOCX = BASE / "Memoria_TFG_DentaLinkLab_Definitiva.docx"

PDF_IMAGES_DIR = BASE / "ilovepdf_images"
STITCH_DIR = BASE / "design_reference" / "stitch_redise" / "stitch_redise_o_de_interfaz_web"


def style_paragraph(paragraph, size=11, bold=False, font="Georgia"):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1):
    style_candidates = {
        1: ["Heading 1", "Encabezado 1", "Titulo 1"],
        2: ["Heading 2", "Encabezado 2", "Titulo 2"],
        3: ["Heading 3", "Encabezado 3", "Titulo 3"],
    }.get(level, ["Heading 1", "Encabezado 1", "Titulo 1"])
    p = None
    for style_name in style_candidates:
        try:
            p = doc.add_paragraph(text, style=style_name)
            break
        except KeyError:
            continue
    if p is None:
        p = doc.add_paragraph(text)
        style_paragraph(p, size=16 if level == 1 else 13, bold=True)
    for run in p.runs:
        run.font.name = "Georgia"
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(11)
    return p


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code.rstrip() + "\n")
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def read_snippet(path: Path, start_marker: str, end_marker: str, max_lines: int = 80) -> str:
    text = path.read_text(encoding="utf-8", errors="ignore")
    start = text.find(start_marker)
    if start == -1:
        return f"# No se encontro el bloque: {start_marker}\n"
    end = text.find(end_marker, start + 1) if end_marker else -1
    if end == -1:
        end = len(text)
    snippet = text[start:end].splitlines()[:max_lines]
    return "\n".join(snippet)


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width_mm: int = 150):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run().add_picture(str(image_path), width=Mm(width_mm))
    c = doc.add_paragraph(caption)
    style_paragraph(c, size=10)


def build():
    if not SOURCE_DOCX.exists():
        raise SystemExit(f"No existe documento base: {SOURCE_DOCX}")
    if not SOURCE_PDF.exists():
        raise SystemExit(f"No existe PDF de referencia: {SOURCE_PDF}")

    doc = Document(str(SOURCE_DOCX))

    doc.add_page_break()
    add_heading(doc, "ANEXO E. Integracion de memoria previa y evidencia tecnica", level=1)
    add_body(
        doc,
        "Este anexo integra los apartados clave de la memoria inicial en PDF y la memoria extendida en DOCX, "
        "consolidando en un unico documento la trazabilidad tecnica del proyecto DentaLinkLab para su defensa final."
    )

    add_heading(doc, "E.1 Modelo de datos (ERD) y entidades criticas", level=2)
    add_body(
        doc,
        "A partir de la memoria original, el modelo de datos se centra en las entidades User, Patient, Order e Invoice. "
        "En la version final se amplia con trazabilidad QR entre clinicas y catalogo unificado de clinicas registradas/publicas."
    )
    for img in sorted(PDF_IMAGES_DIR.glob("*.jpg"))[:2]:
        add_image_with_caption(doc, img, f"Figura E.1 - Evidencia ERD/arquitectura importada desde PDF ({img.name})")

    add_heading(doc, "E.2 Arquitectura de sistema y flujo de trabajo", level=2)
    add_body(
        doc,
        "La arquitectura final mantiene backend Django REST y front desacoplado (React/Vue), con autenticacion JWT, "
        "modulo de pedidos, mensajeria y busqueda de clinicas. El flujo operativo valida estados de pedido, mensajes y trazabilidad."
    )
    for img in sorted(PDF_IMAGES_DIR.glob("*.jpg"))[2:4]:
        add_image_with_caption(doc, img, f"Figura E.2 - Flujo tecnico y arquitectura consolidada ({img.name})")

    add_heading(doc, "E.3 Stack tecnologico, infraestructura y despliegue", level=2)
    add_body(
        doc,
        "Stack principal: Django + DRF, SimpleJWT, PostgreSQL/SQLite en local, frontend en React y migracion progresiva a Vue 3. "
        "Para despliegue se utiliza Docker y entorno cloud con pipeline de actualizacion continua."
    )
    for img in sorted(PDF_IMAGES_DIR.glob("*.jpg"))[4:]:
        add_image_with_caption(doc, img, f"Figura E.3 - Infraestructura/stack desde memoria base ({img.name})")

    add_heading(doc, "E.4 Evidencias visuales combinadas (PDF + Wireframes/Mockups)", level=2)
    add_body(
        doc,
        "En este bloque se combinan evidencias de la memoria original y el rediseño Stitch para demostrar continuidad "
        "entre la propuesta funcional y la implementacion real."
    )
    stitch_samples = [
        STITCH_DIR / "wireframe_buscador_de_cl_nicas_mapa_1" / "screen.png",
        STITCH_DIR / "mockup_buscador_de_cl_nicas_mapa_1" / "screen.png",
        STITCH_DIR / "wireframe_login_recuperar_contrase_a_1" / "screen.png",
        STITCH_DIR / "mockup_login_recuperar_contrase_a_2" / "screen.png",
    ]
    for idx, image in enumerate(stitch_samples, start=1):
        add_image_with_caption(doc, image, f"Figura E.4.{idx} - Propuesta Stitch combinada en memoria final")

    doc.add_page_break()
    add_heading(doc, "ANEXO F. Fragmentos de codigo y APIs clave", level=1)
    add_body(
        doc,
        "Los siguientes bloques recogen codigo real del repositorio para responder preguntas tecnicas del tribunal "
        "sobre mapa, trazabilidad QR y visualizacion STL."
    )

    add_heading(doc, "F.1 API de mapa y filtros del buscador de clinicas", level=2)
    add_body(doc, "Endpoint backend que unifica clinicas registradas y catalogo publico con filtros por texto, precio y rating:")
    add_code_block(
        doc,
        read_snippet(
            BASE / "users" / "views.py",
            "class ClinicDirectoryView",
            "return Response(serializer.data, status=status.HTTP_200_OK)",
            max_lines=95,
        ),
    )

    add_heading(doc, "F.2 API de transferencia de pacientes por QR entre clinicas", level=2)
    add_body(doc, "Acciones share_qr e import_qr implementadas en PatientViewSet:")
    add_code_block(
        doc,
        read_snippet(
            BASE / "marketplace" / "views.py",
            "def share_qr",
            "def get_queryset",
            max_lines=110,
        ),
    )

    add_heading(doc, "F.3 Implementacion del visor STL (Three.js + React Fiber)", level=2)
    add_body(doc, "Componente de visualizacion 3D para escaneados intraorales en formato STL:")
    add_code_block(
        doc,
        read_snippet(
            BASE / "frontend" / "src" / "components" / "Viewer3D.js",
            "function Model",
            "export default Viewer3D;",
            max_lines=110,
        ),
    )

    add_heading(doc, "F.4 Frontend de mapa (Leaflet) y consulta API", level=2)
    add_body(doc, "Integracion en cliente para representacion cartografica y comparador clinico:")
    add_code_block(
        doc,
        read_snippet(
            BASE / "frontend" / "src" / "pages" / "ClinicFinder.js",
            "function ClinicFinder()",
            "export default ClinicFinder;",
            max_lines=120,
        ),
    )

    add_heading(doc, "F.5 Frontend Vue actual (busqueda de clinicas y filtros)", level=2)
    add_body(doc, "Migracion en Vue 3 de la vista Clinic Finder con los mismos parametros de negocio:")
    add_code_block(
        doc,
        read_snippet(
            BASE / "frontend-vue" / "src" / "views" / "ClinicFinderView.vue",
            "<script setup>",
            "</template>",
            max_lines=120,
        ),
    )

    add_heading(doc, "F.6 Instrucciones para indice automatico en Word", level=2)
    add_body(
        doc,
        "Para regenerar el indice en Word: Referencias > Tabla de contenido > Actualizar toda la tabla. "
        "Los titulos de este documento usan estilos de encabezado (Heading 1 / Heading 2), por lo que el indice se actualiza sin edicion manual."
    )

    doc.save(str(OUTPUT_DOCX))
    print(f"Documento definitivo generado: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build()
