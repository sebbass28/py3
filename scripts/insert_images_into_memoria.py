from pathlib import Path
import unicodedata
from collections import defaultdict
from docx import Document
from docx.shared import Mm, Pt


BASE = Path(r"c:\Users\Sebas\Desktop\py3")
DOCX_PATH = BASE / "Memoria_TFG_IES_El_Grao_v2.docx"
STITCH_DIR = BASE / "design_reference" / "stitch_redise" / "stitch_redise_o_de_interfaz_web"
OUTPUT_PATH = BASE / "Memoria_TFG_IES_El_Grao_v2_con_wireframes_mockups.docx"


def humanize(text: str) -> str:
    clean = unicodedata.normalize("NFKD", text)
    clean = "".join(ch for ch in clean if not unicodedata.combining(ch))
    clean = clean.replace("_", " ").strip()
    return " ".join(part.capitalize() for part in clean.split())


def parse_stitch_screens():
    groups = defaultdict(lambda: {"wireframe": [], "mockup": []})
    for screen_path in sorted(STITCH_DIR.glob("*/screen.png")):
        folder_name = screen_path.parent.name
        prefix = "wireframe" if folder_name.startswith("wireframe_") else "mockup" if folder_name.startswith("mockup_") else None
        if not prefix:
            continue
        view_name = folder_name[len(prefix) + 1 :]
        if view_name.endswith("_1") or view_name.endswith("_2"):
            view_name = view_name[:-2]
        groups[view_name][prefix].append(screen_path)
    return groups


def add_heading(paragraph, size: int, bold: bool = True) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    run = paragraph.runs[0]
    run.bold = bold
    run.font.name = "Georgia"
    run.font.size = Pt(size)


def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"No existe el documento base: {DOCX_PATH}")
    if not STITCH_DIR.exists():
        raise SystemExit(f"No existe la carpeta de rediseño Stitch: {STITCH_DIR}")

    grouped = parse_stitch_screens()
    if not grouped:
        raise SystemExit("No se han encontrado pantallas screen.png en el ZIP de Stitch.")

    doc = Document(str(DOCX_PATH))
    doc.add_page_break()

    title = doc.add_paragraph("ANEXO D. WIREFRAMES Y MOCKUPS DEL REDISEÑO")
    add_heading(title, 16)

    intro = doc.add_paragraph(
        "Este anexo incorpora los wireframes y mockups generados en Stitch para todas las vistas funcionales del sistema. "
        "Cada bloque agrupa la propuesta de baja fidelidad (wireframe) y su versión visual final (mockup)."
    )
    for run in intro.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(11)

    figure = 1
    for raw_view_name in sorted(grouped.keys()):
        block = grouped[raw_view_name]
        if not block["wireframe"] and not block["mockup"]:
            continue

        doc.add_page_break()
        section = doc.add_paragraph(f"D.{figure}. {humanize(raw_view_name)}")
        add_heading(section, 13)

        if block["wireframe"]:
            wf_label = doc.add_paragraph("Wireframe")
            add_heading(wf_label, 11)
            for image_path in block["wireframe"]:
                pic = doc.add_paragraph()
                pic.alignment = 1
                pic.add_run().add_picture(str(image_path), width=Mm(160))
                cap = doc.add_paragraph(f"Figura {figure}. Wireframe - {humanize(raw_view_name)}")
                for run in cap.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(10)
                figure += 1

        if block["mockup"]:
            mk_label = doc.add_paragraph("Mockup")
            add_heading(mk_label, 11)
            for image_path in block["mockup"]:
                pic = doc.add_paragraph()
                pic.alignment = 1
                pic.add_run().add_picture(str(image_path), width=Mm(160))
                cap = doc.add_paragraph(f"Figura {figure}. Mockup - {humanize(raw_view_name)}")
                for run in cap.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(10)
                figure += 1

    doc.save(str(OUTPUT_PATH))
    print(f"Documento generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

