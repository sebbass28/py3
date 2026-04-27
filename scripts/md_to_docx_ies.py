from pathlib import Path
import re
from docx import Document
from docx.shared import Mm, Pt


def set_ies_format(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Mm(35)
    section.right_margin = Mm(30)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5


def strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def add_heading(doc: Document, text: str, level: int) -> None:
    # Map markdown levels to readable sizes (max 18 as guide suggests).
    size_map = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}
    p = doc.add_paragraph(strip_inline_md(text.strip()))
    if p.runs:
        run = p.runs[0]
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(size_map.get(level, 11))


def add_paragraph_with_inline(doc: Document, text: str) -> None:
    """
    Parse simple inline markdown tokens:
    - **bold**
    - *italic*
    - `code`
    """
    p = doc.add_paragraph("")
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run("")
        run.font.size = Pt(11)
        run.font.name = "Georgia"
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
        else:
            run.text = part


def is_md_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_separator(line: str) -> bool:
    stripped = line.strip().replace(" ", "")
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    cells = [c for c in stripped.strip("|").split("|")]
    return all(c and set(c) <= {"-", ":"} for c in cells)


def split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_md_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        for line in table_lines:
            add_paragraph_with_inline(doc, line)
        return

    header = split_table_row(table_lines[0])
    separator = table_lines[1]
    body_lines = table_lines[2:] if is_table_separator(separator) else table_lines[1:]

    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for idx, col in enumerate(header):
        table.rows[0].cells[idx].text = strip_inline_md(col)

    for line in body_lines:
        if not is_md_table_line(line):
            continue
        values = split_table_row(line)
        row = table.add_row().cells
        for idx in range(min(len(values), len(header))):
            row[idx].text = strip_inline_md(values[idx])


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_ies_format(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        # Markdown headings: # .. ######
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            if 1 <= level <= 6 and stripped[level:].startswith(" "):
                add_heading(doc, stripped[level + 1 :], level)
                i += 1
                continue

        # Markdown table block
        if is_md_table_line(stripped):
            block = [stripped]
            j = i + 1
            while j < len(lines) and is_md_table_line(lines[j].strip()):
                block.append(lines[j].strip())
                j += 1
            add_md_table(doc, block)
            i = j
            continue

        # Bullet lists
        if stripped.startswith("- "):
            add_paragraph_with_inline(doc, f"• {stripped[2:]}")
            i += 1
            continue

        # Keep numbered lines as-is (e.g., "1. ...")
        add_paragraph_with_inline(doc, stripped)
        i += 1

    doc.save(out_path)


if __name__ == "__main__":
    base = Path(r"c:\Users\Sebas\Desktop\py3")
    convert(
        base / "Memoria_TFG_IES_El_Grao_v2.md",
        base / "Memoria_TFG_IES_El_Grao_v2.docx",
    )
    print(base / "Memoria_TFG_IES_El_Grao_v2.docx")

