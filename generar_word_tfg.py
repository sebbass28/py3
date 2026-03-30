import os
import re
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name, size, bold=False, italic=False):
    run.font.name = font_name
    r = run._element.rPr.get_or_add_rFonts()
    r.set(qn('w:eastAsia'), font_name)
    r.set(qn('w:hAnsi'), font_name)
    r.set(qn('w:ascii'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_formatted_text(paragraph, text, font_name='Georgia', size=11):
    # Regex para detectar negritas **texto** y cursivas *texto*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_font(run, font_name, size, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            run = paragraph.add_run(content)
            set_font(run, font_name, size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, font_name, size)

def create_tfg_word():
    template_path = r'c:\Users\Apliquem\Desktop\py3\PlantillaPortadaProyecto.docx'
    md_path = r'C:\Users\Apliquem\.gemini\antigravity\brain\fec2339a-805a-401d-aef2-15b9df6565bc\memoria_tecnica_tfg.md'
    output_path = r'c:\Users\Apliquem\Desktop\py3\Memoria_Final_DentaLinkLab_DAW.docx'

    if os.path.exists(template_path):
        doc = Document(template_path)
        doc.add_page_break()
    else:
        doc = Document()

    # Márgenes oficiales IES El Grao
    for section in doc.sections:
        section.left_margin = Mm(35)
        section.right_margin = Mm(30)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code = False
    for line in lines:
        raw_line = line.strip()
        
        if raw_line.startswith('---'):
            doc.add_page_break()
            continue

        if raw_line.startswith('```'):
            in_code = not in_code
            continue

        # Encabezados (Manuales para evitar errores de estilos locales)
        if raw_line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(raw_line[2:])
            set_font(run, 'Georgia', 18, bold=True)
            p.paragraph_format.space_after = Pt(20)
            continue
        elif raw_line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(raw_line[3:])
            set_font(run, 'Georgia', 16, bold=True)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(12)
            continue
        elif raw_line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(raw_line[4:])
            set_font(run, 'Georgia', 14, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            continue

        # Párrafos de texto o código
        p = doc.add_paragraph()
        if in_code:
            run = p.add_run(line) # Mantener indentación
            set_font(run, 'Consolas', 10)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        else:
            # Detección de listas manual
            if raw_line.startswith(('- ', '* ')):
                p.paragraph_format.left_indent = Mm(10)
                # Usar un punto medio manual en lugar de estilos automáticos problemáticos
                run_bullet = p.add_run("· ")
                set_font(run_bullet, 'Georgia', 11, bold=True)
                add_formatted_text(p, raw_line[2:])
            elif re.match(r'^\d+\.', raw_line):
                p.paragraph_format.left_indent = Mm(10)
                add_formatted_text(p, raw_line)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                add_formatted_text(p, raw_line)

    try:
        doc.save(output_path)
        print(f"Document updated successfully: {output_path}")
    except PermissionError:
        print(f"Error: El archivo {output_path} está abierto. Ciérrelo e intente de nuevo.")

if __name__ == "__main__":
    create_tfg_word()
